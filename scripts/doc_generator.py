#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Generator - EZTeach Skill
文档生成与导出，完全由脚本控制格式
"""

import os
import json
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentGenerator:
    """文档生成器 - 脚本控制格式生成教学文档"""
    
    def __init__(self):
        self.default_styles = {
            'title': {
                'font_size': 18,
                'bold': True,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                'space_after': 24
            },
            'subtitle': {
                'font_size': 14,
                'bold': True,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                'space_after': 18
            },
            'heading1': {
                'font_size': 16,
                'bold': True,
                'space_before': 18,
                'space_after': 12
            },
            'heading2': {
                'font_size': 14,
                'bold': True,
                'space_before': 12,
                'space_after': 6
            },
            'heading3': {
                'font_size': 12,
                'bold': True,
                'space_before': 6,
                'space_after': 6
            },
            'normal': {
                'font_size': 12,
                'line_spacing': 1.5,
                'first_line_indent': 24
            },
            'list': {
                'font_size': 12,
                'left_indent': 24,
                'space_after': 3
            },
            'table_header': {
                'font_size': 11,
                'bold': True,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER
            },
            'table_cell': {
                'font_size': 11,
                'alignment': WD_ALIGN_PARAGRAPH.LEFT
            },
            'info': {
                'font_size': 10.5,
                'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
        }
        self.doc = None
    
    def create_new_document(self, template_path: Optional[str] = None) -> None:
        """
        创建新文档
        
        Args:
            template_path: 可选模板路径
        """
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
        else:
            self.doc = Document()
            self._setup_default_styles()
    
    def _setup_default_styles(self) -> None:
        """设置默认文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
        
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
    
    def _apply_style(self, para, style_name: str) -> None:
        """应用样式到段落"""
        style = self.default_styles.get(style_name, self.default_styles['normal'])
        
        for run in para.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(style.get('font_size', 12))
            run.font.bold = style.get('bold', False)
        
        para.alignment = style.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if 'space_before' in style:
            para.paragraph_format.space_before = Pt(style['space_before'])
        if 'space_after' in style:
            para.paragraph_format.space_after = Pt(style['space_after'])
        if 'line_spacing' in style:
            para.paragraph_format.line_spacing = style['line_spacing']
        if 'first_line_indent' in style:
            para.paragraph_format.first_line_indent = Pt(style['first_line_indent'])
        if 'left_indent' in style:
            para.paragraph_format.left_indent = Pt(style['left_indent'])
    
    def add_title(self, text: str) -> None:
        """添加标题"""
        para = self.doc.add_paragraph(text)
        self._apply_style(para, 'title')
    
    def add_subtitle(self, text: str) -> None:
        """添加副标题"""
        para = self.doc.add_paragraph(text)
        self._apply_style(para, 'subtitle')
    
    def add_heading(self, text: str, level: int = 1) -> None:
        """添加标题"""
        style_map = {1: 'heading1', 2: 'heading2', 3: 'heading3'}
        style_name = style_map.get(level, 'heading1')
        para = self.doc.add_paragraph(text)
        self._apply_style(para, style_name)
    
    def add_paragraph(self, text: str, style: str = 'normal') -> None:
        """添加段落"""
        para = self.doc.add_paragraph(text)
        self._apply_style(para, style)
    
    def add_list(self, items: List[str], ordered: bool = False) -> None:
        """
        添加列表
        
        Args:
            items: 列表项
            ordered: 是否有序列表
        """
        for i, item in enumerate(items):
            prefix = f"{i+1}. " if ordered else "● "
            para = self.doc.add_paragraph(prefix + item)
            self._apply_style(para, 'list')
    
    def add_table(self, data: List[List[str]], headers: Optional[List[str]] = None,
                  col_widths: Optional[List[float]] = None) -> None:
        """
        添加表格
        
        Args:
            data: 表格数据
            headers: 表头
            col_widths: 列宽（厘米）
        """
        rows = len(data) + (1 if headers else 0)
        cols = len(headers) if headers else len(data[0]) if data else 0
        
        if cols == 0:
            return
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Cm(width)
        
        row_idx = 0
        
        # 添加表头
        if headers:
            for col_idx, header in enumerate(headers):
                cell = table.cell(row_idx, col_idx)
                cell.text = header
                for para in cell.paragraphs:
                    self._apply_style(para, 'table_header')
            row_idx += 1
        
        # 添加数据
        for row_data in data:
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(cell_data)
                    for para in cell.paragraphs:
                        self._apply_style(para, 'table_cell')
            row_idx += 1
    
    def add_info_row(self, label: str, value: str) -> None:
        """添加信息行（如：学科：语文）"""
        para = self.doc.add_paragraph()
        run_label = para.add_run(f"{label}：")
        run_label.font.bold = True
        run_label.font.size = Pt(10.5)
        run_value = para.add_run(value)
        run_value.font.size = Pt(10.5)
        para.paragraph_format.space_after = Pt(3)
    
    def add_page_break(self) -> None:
        """添加分页符"""
        self.doc.add_page_break()
    
    def generate_teaching_plan(self, plan_data: Dict[str, Any], output_path: str) -> str:
        """
        生成教案文档
        
        Args:
            plan_data: 教案数据
            output_path: 输出路径
            
        Returns:
            输出文件路径
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装")
        
        self.create_new_document()
        
        # 标题
        self.add_title(plan_data.get('title', '教案'))
        
        # 基本信息表格
        basic_info = plan_data.get('basic_info', {})
        info_data = [
            ['学科', basic_info.get('subject', ''), '年级', basic_info.get('grade', '')],
            ['课时', basic_info.get('period', ''), '授课教师', basic_info.get('teacher', '')],
            ['授课时间', basic_info.get('date', datetime.now().strftime('%Y年%m月%d日')), '课型', basic_info.get('type', '新授课')]
        ]
        self.add_table(info_data, col_widths=[2, 4, 2, 4])
        
        self.doc.add_paragraph()  # 空行
        
        # 教学目标
        if 'teaching_objectives' in plan_data:
            self.add_heading('一、教学目标', 1)
            objectives = plan_data['teaching_objectives']
            if isinstance(objectives, list):
                self.add_list(objectives)
            else:
                self.add_paragraph(str(objectives))
        
        # 教学重点
        if 'key_points' in plan_data:
            self.add_heading('二、教学重点', 1)
            self.add_paragraph(plan_data['key_points'])
        
        # 教学难点
        if 'difficult_points' in plan_data:
            self.add_heading('三、教学难点', 1)
            self.add_paragraph(plan_data['difficult_points'])
        
        # 教学准备
        if 'preparation' in plan_data:
            self.add_heading('四、教学准备', 1)
            self.add_paragraph(plan_data['preparation'])
        
        # 教学过程
        if 'teaching_process' in plan_data:
            self.add_heading('五、教学过程', 1)
            process_steps = plan_data['teaching_process']
            if isinstance(process_steps, list):
                for i, step in enumerate(process_steps, 1):
                    if isinstance(step, dict):
                        self.add_heading(f"{i}. {step.get('title', '')}", 2)
                        self.add_paragraph(step.get('content', ''))
                        if 'activity' in step:
                            self.add_paragraph(f"活动设计：{step['activity']}")
                        if 'time' in step:
                            self.add_paragraph(f"时间安排：{step['time']}")
                    else:
                        self.add_paragraph(f"{i}. {step}")
        
        # 板书设计
        if 'blackboard_design' in plan_data:
            self.add_heading('六、板书设计', 1)
            self.add_paragraph(plan_data['blackboard_design'])
        
        # 作业布置
        if 'homework' in plan_data:
            self.add_heading('七、作业布置', 1)
            homework = plan_data['homework']
            if isinstance(homework, list):
                self.add_list(homework)
            else:
                self.add_paragraph(str(homework))
        
        # 教学反思
        if 'reflection' in plan_data:
            self.add_heading('八、教学反思', 1)
            self.add_paragraph(plan_data['reflection'])
        
        # 保存
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        
        return str(output_path)
    
    def generate_learning_sheet(self, sheet_data: Dict[str, Any], output_path: str) -> str:
        """
        生成学习单文档
        
        Args:
            sheet_data: 学习单数据
            output_path: 输出路径
            
        Returns:
            输出文件路径
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装")
        
        self.create_new_document()
        
        # 标题
        self.add_title(sheet_data.get('title', '学习单'))
        
        # 基本信息
        self.add_info_row('班级', sheet_data.get('class', ''))
        self.add_info_row('姓名', sheet_data.get('name', ''))
        self.add_info_row('学号', sheet_data.get('student_id', ''))
        
        self.doc.add_paragraph()
        
        # 学习目标
        if 'learning_objectives' in sheet_data:
            self.add_heading('一、学习目标', 1)
            self.add_list(sheet_data['learning_objectives'])
        
        # 预习内容
        if 'preview' in sheet_data:
            self.add_heading('二、预习内容', 1)
            preview = sheet_data['preview']
            if isinstance(preview, list):
                self.add_list(preview)
            else:
                self.add_paragraph(str(preview))
        
        # 课堂练习
        if 'exercises' in sheet_data:
            self.add_heading('三、课堂练习', 1)
            exercises = sheet_data['exercises']
            for i, exercise in enumerate(exercises, 1):
                self.add_paragraph(f"{i}. {exercise}")
                self.doc.add_paragraph()
                self.doc.add_paragraph()
        
        # 拓展思考
        if 'thinking' in sheet_data:
            self.add_heading('四、拓展思考', 1)
            self.add_paragraph(sheet_data['thinking'])
        
        # 学习收获
        self.add_heading('五、学习收获', 1)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 保存
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        
        return str(output_path)
    
    def batch_generate(self, data_list: List[Dict[str, Any]], 
                       output_dir: str, doc_type: str = 'plan') -> Dict[str, Any]:
        """
        批量生成文档
        
        Args:
            data_list: 数据列表
            output_dir: 输出目录
            doc_type: 文档类型 'plan' 或 'sheet'
            
        Returns:
            生成结果
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        results = {
            "success": 0,
            "failed": 0,
            "files": []
        }
        
        generator_method = self.generate_teaching_plan if doc_type == 'plan' else self.generate_learning_sheet
        
        for i, data in enumerate(data_list):
            try:
                filename = f"{doc_type}_{i+1}.docx"
                output_path = output_dir / filename
                generator_method(data, str(output_path))
                results["success"] += 1
                results["files"].append(str(output_path))
            except Exception as e:
                results["failed"] += 1
                results["files"].append(f"Error: {str(e)}")
        
        return results


def generate_document(data_path: str, output_path: str, doc_type: str = 'plan') -> Dict[str, Any]:
    """
    生成文档主函数
    
    Args:
        data_path: JSON数据文件路径
        output_path: 输出文件路径
        doc_type: 文档类型
        
    Returns:
        生成结果
    """
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    generator = DocumentGenerator()
    result = {
        "success": True,
        "data_file": data_path,
        "output_file": output_path,
        "doc_type": doc_type
    }
    
    try:
        if doc_type == 'plan':
            generator.generate_teaching_plan(data, output_path)
        elif doc_type == 'sheet':
            generator.generate_learning_sheet(data, output_path)
        result["generated"] = True
    except Exception as e:
        result["success"] = False
        result["error"] = str(e)
    
    return result


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) >= 3:
        data_file = sys.argv[1]
        output_file = sys.argv[2]
        doc_type = sys.argv[3] if len(sys.argv) > 3 else 'plan'
        result = generate_document(data_file, output_file, doc_type)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("Usage: python doc_generator.py <data.json> <output.docx> [plan|sheet]")
