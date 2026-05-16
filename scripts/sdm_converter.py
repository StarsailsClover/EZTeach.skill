#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SDM Converter - EZTeach Skill
Semantic Document Markdown Converter
Word↔Markdown双向无损转换、格式语义捕获与复用
"""

import os
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
from io import StringIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    from markdown.extensions.toc import TocExtension
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class SDMConverter:
    """语义文档Markdown转换器"""
    
    def __init__(self):
        self.style_mapping = {
            'heading 1': {'level': 1, 'markdown': '#'},
            'heading 2': {'level': 2, 'markdown': '##'},
            'heading 3': {'level': 3, 'markdown': '###'},
            'heading 4': {'level': 4, 'markdown': '####'},
            'heading 5': {'level': 5, 'markdown': '#####'},
            'heading 6': {'level': 6, 'markdown': '######'},
            'list paragraph': {'list': True},
            'normal': {'normal': True},
        }
        
        # 标题正则
        self.heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
        # 列表正则
        self.list_pattern = re.compile(r'^(\s*)([-*+]|\d+\.)\s+(.+)$')
        # 粗体正则
        self.bold_pattern = re.compile(r'\*\*(.+?)\*\*')
        # 斜体正则
        self.italic_pattern = re.compile(r'\*(.+?)\*')
        # 代码正则
        self.code_pattern = re.compile(r'`(.+?)`')
    
    def docx_to_markdown(self, docx_path: str, output_path: Optional[str] = None) -> str:
        """
        Word文档转换为Markdown
        
        Args:
            docx_path: Word文档路径
            output_path: 输出Markdown路径
            
        Returns:
            转换后的Markdown内容
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请执行: pip install python-docx")
        
        doc = Document(docx_path)
        md_lines = []
        semantic_styles = []
        
        for para in doc.paragraphs:
            style_name = para.style.name.lower()
            text = para.text.strip()
            
            if not text:
                md_lines.append('')
                continue
            
            # 记录语义样式
            style_info = {
                'text': text,
                'style': style_name,
                'alignment': para.alignment,
                'runs': []
            }
            
            # 提取run格式信息
            for run in para.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_name': run.font.name
                }
                style_info['runs'].append(run_info)
            semantic_styles.append(style_info)
            
            # 转换为Markdown
            md_text = text
            
            # 处理粗体和斜体
            for run in para.runs:
                if run.bold and run.text:
                    md_text = md_text.replace(run.text, f'**{run.text}**')
                elif run.italic and run.text:
                    md_text = md_text.replace(run.text, f'*{run.text}*')
            
            # 处理标题
            if 'heading' in style_name:
                level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                prefix = '#' * min(level, 6)
                md_lines.append(f'{prefix} {md_text}')
            # 处理列表
            elif para.style.name == 'List Paragraph' or para.left_indent:
                md_lines.append(f'- {md_text}')
            # 普通段落
            else:
                md_lines.append(md_text)
        
        md_content = '\n\n'.join(md_lines)
        
        # 添加语义样式元数据
        md_content = f'<!-- SEMANTIC_STYLES: {len(semantic_styles)} -->\n\n' + md_content
        
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
        
        return md_content
    
    def markdown_to_docx(self, markdown_path: str, output_path: Optional[str] = None, 
                         template_path: Optional[str] = None) -> str:
        """
        Markdown转换为Word文档
        
        Args:
            markdown_path: Markdown文件路径
            output_path: 输出Word路径
            template_path: 模板文档路径（用于样式复用）
            
        Returns:
            输出文件路径
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请执行: pip install python-docx")
        
        # 读取Markdown内容
        with open(markdown_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 使用模板或新建文档
        if template_path and Path(template_path).exists():
            doc = Document(template_path)
            # 清空模板内容但保留样式
            for para in doc.paragraphs:
                para.clear()
        else:
            doc = Document()
        
        lines = md_content.split('\n')
        in_code_block = False
        code_lines = []
        
        for line in lines:
            # 跳过元数据注释
            if line.startswith('<!--') and line.endswith('-->'):
                continue
            
            # 代码块处理
            if line.startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block and code_lines:
                    # 添加代码块
                    para = doc.add_paragraph()
                    code_para = doc.add_paragraph('\n'.join(code_lines))
                    code_para.style = 'Quote'
                    code_lines = []
                continue
            
            if in_code_block:
                code_lines.append(line)
                continue
            
            if not line.strip():
                continue
            
            # 处理标题
            heading_match = self.heading_pattern.match(line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                doc.add_heading(self._process_inline_format(text), level=level)
                continue
            
            # 处理列表
            list_match = self.list_pattern.match(line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(3)
                para = doc.add_paragraph(self._process_inline_format(text), style='List Paragraph')
                para.paragraph_format.left_indent = Pt(indent * 20)
                continue
            
            # 普通段落
            para = doc.add_paragraph()
            self._add_formatted_text(para, line)
        
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            return str(output_path)
        
        return doc
    
    def _process_inline_format(self, text: str) -> str:
        """处理内联格式（移除Markdown标记）"""
        text = self.bold_pattern.sub(r'\1', text)
        text = self.italic_pattern.sub(r'\1', text)
        text = self.code_pattern.sub(r'\1', text)
        return text
    
    def _add_formatted_text(self, paragraph, text: str):
        """添加带格式的文本到段落"""
        # 简单处理：先处理粗体
        parts = self.bold_pattern.split(text)
        for i, part in enumerate(parts):
            if i % 2 == 1:
                run = paragraph.add_run(part)
                run.bold = True
            else:
                # 处理斜体
                italic_parts = self.italic_pattern.split(part)
                for j, ipart in enumerate(italic_parts):
                    if j % 2 == 1:
                        run = paragraph.add_run(ipart)
                        run.italic = True
                    else:
                        paragraph.add_run(ipart)
    
    def capture_format_semantics(self, docx_path: str) -> Dict[str, Any]:
        """
        捕获文档格式语义
        
        Args:
            docx_path: Word文档路径
            
        Returns:
            格式语义字典
        """
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not available"}
        
        doc = Document(docx_path)
        semantics = {
            "styles": {},
            "paragraphs": [],
            "fonts": set(),
            "colors": set()
        }
        
        # 提取所有样式
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_info = {
                    "name": style.name,
                    "font_name": style.font.name,
                    "font_size": style.font.size.pt if style.font.size else None,
                    "bold": style.font.bold,
                    "italic": style.font.italic,
                    "alignment": style.paragraph_format.alignment
                }
                semantics["styles"][style.name] = style_info
        
        # 提取段落语义
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            
            para_info = {
                "index": i,
                "text": para.text.strip(),
                "style": para.style.name,
                "alignment": str(para.alignment),
                "level": self._detect_heading_level(para),
                "is_list": self._is_list_item(para)
            }
            semantics["paragraphs"].append(para_info)
            
            for run in para.runs:
                if run.font.name:
                    semantics["fonts"].add(run.font.name)
                if run.font.color and run.font.color.rgb:
                    semantics["colors"].add(str(run.font.color.rgb))
        
        semantics["fonts"] = list(semantics["fonts"])
        semantics["colors"] = list(semantics["colors"])
        
        return semantics
    
    def _detect_heading_level(self, para) -> int:
        """检测标题级别"""
        style_name = para.style.name.lower()
        if 'heading' in style_name:
            parts = style_name.split()
            if len(parts) > 1 and parts[-1].isdigit():
                return int(parts[-1])
        return 0
    
    def _is_list_item(self, para) -> bool:
        """检测是否为列表项"""
        return (para.style.name == 'List Paragraph' or 
                para.paragraph_format.left_indent is not None)
    
    def batch_convert(self, input_dir: str, output_dir: str, direction: str = 'docx2md') -> Dict[str, Any]:
        """
        批量转换
        
        Args:
            input_dir: 输入目录
            output_dir: 输出目录
            direction: 转换方向 'docx2md' 或 'md2docx'
            
        Returns:
            转换结果
        """
        input_dir = Path(input_dir)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        results = {
            "success": 0,
            "failed": 0,
            "files": []
        }
        
        if direction == 'docx2md':
            pattern = '*.docx'
            converter = self.docx_to_markdown
            ext = '.md'
        else:
            pattern = '*.md'
            converter = self.markdown_to_docx
            ext = '.docx'
        
        for file_path in input_dir.glob(pattern):
            try:
                output_path = output_dir / f"{file_path.stem}{ext}"
                converter(str(file_path), str(output_path))
                results["success"] += 1
                results["files"].append({
                    "input": str(file_path),
                    "output": str(output_path),
                    "status": "success"
                })
            except Exception as e:
                results["failed"] += 1
                results["files"].append({
                    "input": str(file_path),
                    "error": str(e),
                    "status": "failed"
                })
        
        return results


def convert_file(input_path: str, output_path: Optional[str] = None, 
                 direction: Optional[str] = None) -> Dict[str, Any]:
    """
    转换文件主函数
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        direction: 转换方向（自动检测）
        
    Returns:
        转换结果
    """
    converter = SDMConverter()
    input_path = Path(input_path)
    
    # 自动检测转换方向
    if direction is None:
        if input_path.suffix.lower() == '.docx':
            direction = 'docx2md'
        elif input_path.suffix.lower() == '.md':
            direction = 'md2docx'
        else:
            return {"error": "无法自动检测转换方向，请指定direction参数"}
    
    result = {
        "input": str(input_path),
        "direction": direction,
        "success": True
    }
    
    try:
        if direction == 'docx2md':
            md_content = converter.docx_to_markdown(str(input_path), output_path)
            result["output"] = output_path or "inline"
            result["content_length"] = len(md_content)
        else:
            docx_path = converter.markdown_to_docx(str(input_path), output_path)
            result["output"] = docx_path
    except Exception as e:
        result["success"] = False
        result["error"] = str(e)
    
    return result


if __name__ == "__main__":
    import sys
    import json
    
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        result = convert_file(input_file, output_file)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("Usage: python sdm_converter.py <input_file> [output_file]")
