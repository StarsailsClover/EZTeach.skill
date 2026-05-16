#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Template Adapter - EZTeach Skill
模板结构解析、语义映射、格式100%复用
"""

import os
import re
import json
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TemplateAdapter:
    """模板适配器 - 解析模板结构并实现内容映射"""
    
    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path
        self.template_structure = None
        self.style_registry = {}
        self.placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
        self.section_markers = {
            'teaching_objectives': ['教学目标', '学习目标', '知识目标', '能力目标'],
            'key_points': ['教学重点', '重点', '核心内容'],
            'difficult_points': ['教学难点', '难点', '易错点'],
            'teaching_process': ['教学过程', '教学流程', '教学步骤'],
            'blackboard_design': ['板书设计', '板书'],
            'homework': ['作业布置', '作业', '课后练习'],
            'reflection': ['教学反思', '反思', '课后反思']
        }
    
    def parse_template_structure(self, template_path: Optional[str] = None) -> Dict[str, Any]:
        """
        解析Word模板结构
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            模板结构字典
        """
        path = template_path or self.template_path
        if not path or not Path(path).exists():
            return {"error": "模板文件不存在"}
        
        if not DOCX_AVAILABLE:
            return {"error": "python-docx库未安装"}
        
        doc = Document(path)
        structure = {
            "file_name": Path(path).name,
            "sections": [],
            "placeholders": [],
            "styles": {},
            "tables": [],
            "paragraph_count": len(doc.paragraphs)
        }
        
        # 解析样式
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                self.style_registry[style.name] = {
                    "font": style.font.name,
                    "font_size": style.font.size.pt if style.font.size else None,
                    "bold": style.font.bold,
                    "italic": style.font.italic,
                    "alignment": style.paragraph_format.alignment,
                    "space_before": style.paragraph_format.space_before.pt if style.paragraph_format.space_before else 0,
                    "space_after": style.paragraph_format.space_after.pt if style.paragraph_format.space_after else 0
                }
        structure["styles"] = self.style_registry
        
        # 解析段落和占位符
        current_section = None
        section_content = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # 检测占位符
            placeholders = self.placeholder_pattern.findall(text)
            for ph in placeholders:
                structure["placeholders"].append({
                    "name": ph,
                    "paragraph_index": i,
                    "full_text": text,
                    "style": para.style.name
                })
            
            # 检测章节
            section_type = self._detect_section_type(text)
            if section_type:
                if current_section:
                    structure["sections"].append({
                        "type": current_section,
                        "content": section_content,
                        "start_index": i - len(section_content)
                    })
                current_section = section_type
                section_content = [{"text": text, "style": para.style.name, "index": i}]
            elif current_section:
                section_content.append({"text": text, "style": para.style.name, "index": i})
        
        # 添加最后一个章节
        if current_section:
            structure["sections"].append({
                "type": current_section,
                "content": section_content
            })
        
        # 解析表格
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": []
            }
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    table_info["cells"].append({
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell.text.strip(),
                        "placeholders": self.placeholder_pattern.findall(cell.text)
                    })
            structure["tables"].append(table_info)
        
        self.template_structure = structure
        return structure
    
    def _detect_section_type(self, text: str) -> Optional[str]:
        """检测文本对应的章节类型"""
        text_lower = text.lower()
        for section_type, markers in self.section_markers.items():
            for marker in markers:
                if marker in text:
                    return section_type
        return None
    
    def extract_placeholder_schema(self) -> Dict[str, Any]:
        """
        提取占位符Schema
        
        Returns:
            占位符Schema字典
        """
        if not self.template_structure:
            self.parse_template_structure()
        
        schema = {
            "placeholders": {},
            "required": [],
            "optional": []
        }
        
        for ph in self.template_structure.get("placeholders", []):
            ph_name = ph["name"]
            schema["placeholders"][ph_name] = {
                "style": ph["style"],
                "location": ph["paragraph_index"],
                "context": ph["full_text"]
            }
            
            # 判断是否必填
            if any(key in ph_name.lower() for key in ['title', 'subject', 'grade', 'content']):
                schema["required"].append(ph_name)
            else:
                schema["optional"].append(ph_name)
        
        return schema
    
    def map_content_to_template(self, content_data: Dict[str, Any], 
                                output_path: str) -> str:
        """
        将内容数据映射到模板并生成新文档
        
        Args:
            content_data: 内容数据字典
            output_path: 输出文件路径
            
        Returns:
            输出文件路径
        """
        if not self.template_path or not Path(self.template_path).exists():
            raise ValueError("模板路径未设置或不存在")
        
        doc = Document(self.template_path)
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            self._replace_placeholders_in_paragraph(para, content_data)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(para, content_data)
        
        # 保存文档
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _replace_placeholders_in_paragraph(self, para, content_data: Dict[str, Any]):
        """替换段落中的占位符（保留格式）"""
        placeholders = self.placeholder_pattern.findall(para.text)
        
        for ph_name in placeholders:
            placeholder = f'{{{{{ph_name}}}}}'
            replacement = str(content_data.get(ph_name, ''))
            
            if placeholder in para.text:
                # 简单替换（保留整体样式）
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
                    elif ph_name in run.text:
                        run.text = run.text.replace(ph_name, replacement)
    
    def clone_format_from_template(self, source_doc_path: str, target_doc_path: str,
                                   output_path: str) -> str:
        """
        从模板克隆格式到目标文档
        
        Args:
            source_doc_path: 源模板文档
            target_doc_path: 目标内容文档
            output_path: 输出路径
            
        Returns:
            输出文件路径
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required")
        
        # 打开源模板获取样式
        source_doc = Document(source_doc_path)
        target_doc = Document(target_doc_path)
        
        # 创建新文档使用模板样式
        result_doc = Document(source_doc_path)
        
        # 清空结果文档内容
        for para in result_doc.paragraphs:
            para.clear()
        
        # 复制目标内容并应用模板样式
        for i, para in enumerate(target_doc.paragraphs):
            if not para.text.strip():
                result_doc.add_paragraph()
                continue
            
            # 匹配最合适的样式
            matched_style = self._match_style(para, source_doc)
            
            new_para = result_doc.add_paragraph(style=matched_style)
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        result_doc.save(str(output_path))
        
        return str(output_path)
    
    def _match_style(self, para, source_doc) -> str:
        """匹配合适的样式"""
        # 简单的样式匹配逻辑
        if para.style.name.startswith('Heading'):
            return para.style.name
        
        # 检测是否为标题
        for run in para.runs:
            if run.bold and run.font.size and run.font.size > Pt(14):
                return 'Heading 2'
        
        return 'Normal'
    
    def generate_mapping_report(self) -> Dict[str, Any]:
        """
        生成模板映射报告
        
        Returns:
            映射报告字典
        """
        if not self.template_structure:
            self.parse_template_structure()
        
        report = {
            "template_name": self.template_structure.get("file_name"),
            "structure_summary": {
                "total_paragraphs": self.template_structure.get("paragraph_count", 0),
                "sections_found": len(self.template_structure.get("sections", [])),
                "placeholders_found": len(self.template_structure.get("placeholders", [])),
                "tables_found": len(self.template_structure.get("tables", []))
            },
            "sections": [s["type"] for s in self.template_structure.get("sections", [])],
            "placeholder_list": [p["name"] for p in self.template_structure.get("placeholders", [])],
            "style_count": len(self.template_structure.get("styles", {}))
        }
        
        return report
    
    def export_template_schema(self, output_path: str) -> str:
        """
        导出模板Schema为JSON
        
        Args:
            output_path: 输出JSON路径
            
        Returns:
            输出文件路径
        """
        schema = {
            "structure": self.template_structure,
            "placeholder_schema": self.extract_placeholder_schema(),
            "mapping_report": self.generate_mapping_report()
        }
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(schema, f, ensure_ascii=False, indent=2)
        
        return str(output_path)


def adapt_template(template_path: str, content_data: Dict[str, Any], 
                   output_path: str) -> Dict[str, Any]:
    """
    模板适配主函数
    
    Args:
        template_path: 模板文件路径
        content_data: 内容数据
        output_path: 输出文件路径
        
    Returns:
        适配结果
    """
    adapter = TemplateAdapter(template_path)
    structure = adapter.parse_template_structure()
    
    result = {
        "success": True,
        "template": template_path,
        "output": output_path,
        "structure_analyzed": False,
        "placeholders_replaced": 0
    }
    
    if "error" not in structure:
        result["structure_analyzed"] = True
        result["sections_found"] = len(structure.get("sections", []))
        result["placeholders_found"] = len(structure.get("placeholders", []))
    
    try:
        adapter.map_content_to_template(content_data, output_path)
        result["placeholders_replaced"] = len(structure.get("placeholders", []))
    except Exception as e:
        result["success"] = False
        result["error"] = str(e)
    
    return result


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) >= 4:
        template = sys.argv[1]
        content_json = sys.argv[2]
        output = sys.argv[3]
        
        with open(content_json, 'r', encoding='utf-8') as f:
            content = json.load(f)
        
        result = adapt_template(template, content, output)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("Usage: python template_adapter.py <template.docx> <content.json> <output.docx>")
